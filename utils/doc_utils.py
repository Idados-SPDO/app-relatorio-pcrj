from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
import pandas as pd


def add_header_paragraphs(doc: Document, validade: str) -> None:
    """
    Adiciona os 5 parágrafos de cabeçalho fixo (Secretaria, link, título, explicação e validade) 
    em um Document do python-docx.
    """
    section = doc.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # 1) SECRETARIA DE EDUCAÇÃO, em negrito
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run("SECRETARIA DE EDUCAÇÃO")
    run1.bold = True
    run1.font.name = "Arial"
    run1.font.size = Pt(10)

    # 2) Link azul
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run("http://www.rio.rj.gov.br/web/sme/pnae")
    run2.underline = True
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0, 0, 255)

    # 3) Título da tabela
    p3 = doc.add_paragraph("Tabela de Preços de Mercado de Gêneros Alimentícios")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_before = Pt(0)
    run3 = p3.runs[0]
    run3.font.name = "Arial"
    run3.font.size = Pt(10)

    # 4) Texto explicativo
    p4 = doc.add_paragraph(
        "A tabela é referência para as aquisições realizadas pelos diversos órgãos do município "
        "e tem o preço dos itens apurado conforme estabelecido no Art. 1º do Decreto nº 51.017/2022 "
        "e alterações, que estabelece que o preço praticado pelo município e divulgado nesta tabela "
        "seja um preço intermediário entre os preços no mercado de atacado e de varejo."
    )
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.line_spacing = 1
    for run in p4.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)

    # 5) Validade (dinâmica)
    p5 = doc.add_paragraph(f"Validade: {validade}")
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.line_spacing = 1


def generate_full_doc(df: pd.DataFrame, validade: str) -> bytes:
    """
    Gera um .docx com todas as colunas (6): 
    ["Código do Item", "Descrição do Item", "Unidade", "Preço Atacado", "Preço Varejo", "Preço Praticado"].
    Retorna os bytes do documento.
    """
    doc = Document()
    add_header_paragraphs(doc, validade)

    quartil_doc = df.copy()

    def combine_desc(row):
        prod = str(row["Produto"])
        desc = row["Descrição"]
        if pd.isna(desc):
            return prod
        desc_str = str(desc).strip()
        if not desc_str or desc_str == "-":
            return prod
        return f"{prod}\n{desc_str}"

    quartil_doc["Descrição do Item"] = quartil_doc.apply(combine_desc, axis=1)

    cols = [
        "Código do Item",
        "Descrição do Item",
        "Unidade",
        "Preço Atacado",
        "Preço Varejo",
        "Preço Praticado",
    ]
    quartil_doc = quartil_doc[cols]

    rows, cols_count = quartil_doc.shape
    table = doc.add_table(rows=rows + 1, cols=cols_count)
    table.style = "Table Grid"
    table.allow_autofit = False

    col_widths = [Cm(5), Cm(15), Cm(2), Cm(2), Cm(2), Cm(2)]
    for idx, width in enumerate(col_widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    # Cabeçalhos
    for j, cell in enumerate(table.rows[0].cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        cell.height = Pt(20)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(quartil_doc.columns[j])
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)

    # Linhas de dados
    for i, row in enumerate(quartil_doc.itertuples(index=False), start=1):
        table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[i].height = Pt(18)
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing = 1
            if j == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(value))
            run.font.name = "Arial"
            run.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_price_only_doc(df: pd.DataFrame, validade: str) -> bytes:
    """
    Gera um .docx apenas com as 4 colunas + coluna "Nº": 
    ["Nº", "Código do Item", "Descrição do Item", "Unidade", "Preço (em R$)"].
    Retorna os bytes do documento.
    """
    doc_price = Document()
    add_header_paragraphs(doc_price, validade)

    price_df = df.copy()

    def combine_desc(r):
        prod = str(r["Produto"])
        d = r["Descrição"]
        if pd.isna(d) or str(d).strip() == "-":
            return prod
        return prod + "\n" + str(d).strip()

    price_df["Descrição do Item"] = price_df.apply(combine_desc, axis=1)
    price_df["Preço (em R$)"] = price_df["Preço Praticado"].apply(
        lambda x: f"{x:.2f}".replace(".", ",") if pd.notna(x) else ""
    )

    cols = ["Código do Item", "Descrição do Item", "Unidade", "Preço (em R$)"]
    tbl_df = price_df[cols].copy()
    tbl_df.insert(0, "Nº", range(1, len(tbl_df) + 1))

    table = doc_price.add_table(rows=tbl_df.shape[0] + 1, cols=tbl_df.shape[1])
    table.style = "Table Grid"
    table.allow_autofit = False

    col_widths = [Cm(1.0), Cm(5), Cm(15), Cm(2), Cm(5)]
    for idx, width in enumerate(col_widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    # Cabeçalhos
    for j, h in enumerate(tbl_df.columns):
        cell = table.cell(0, j)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.height = Pt(20)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)

    # Dados
    for i, row in enumerate(tbl_df.itertuples(index=False), start=1):
        table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[i].height = Pt(18)
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.line_spacing = 1
            # "Descrição do Item" agora está no índice 2
            if j == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(8)

    buf = BytesIO()
    doc_price.save(buf)
    buf.seek(0)
    return buf.getvalue()
